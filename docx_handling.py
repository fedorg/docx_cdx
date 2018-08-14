default_relpath = 'word/_rels/document.xml.rels'

def read_zip_files(filename, paths,*, errors='ignore'):
    from zipfile import ZipFile
    with ZipFile(filename) as myzip:
        for p in paths:
            try:
                with myzip.open(p) as f:
                    yield f.read()
            except Exception as e:
                if errors != 'ignore':
                    raise
                else:
                    yield None
    return


def get_embedding_rels(filename, relpath=None):
    if relpath is None:
        relpath = default_relpath
    from lxml.etree import parse
    from io import BytesIO
    from pathlib import PurePath as P
    cnt = [*read_zip_files(filename, [relpath], errors='raise')][0]
    def rejoin(path):
        ret = []
        for a in path.parts:
            if a == '..': ret.pop()
            else: ret.append(a)
        return str(P(*ret))
    
    tree = parse(BytesIO(cnt))
    base = P(relpath).parent.parent
    def retarget(d):
        return {**d, 'Target': rejoin(base.joinpath(d['Target']))} if 'Target' in d else d
    rels = [retarget(dict(r.items())) for r in tree.findall('Relationship', tree.getroot().nsmap)]
    return rels


def rels_to_paths(rels, filter=None):
    if filter is None:
        filter = lambda rel: rel.get('Type', '').endswith('oleObject')
    paths = [(rel['Id'], rel['Target']) for rel in rels if filter(rel)]
    return paths


def get_embedding_format(compobj):
    from struct import unpack
    b = bytes(compobj)
    start = 28
    ret = []
    for i in range(3):
        try:
            strl, = unpack('<I', b[start:start + 4])
            s, = unpack(f'{strl}s', b[start + 4:start + 4 + strl])
            start = start + 4 + strl
            ret.append(s[:-1])
        except:
            break
    return ret


def read_ole_contents(oledata, paths=None):
    if paths is None:
        paths = [['CONTENTS'], ['\x03PRINT']]
    import olefile
    def try_default(func, *args, **kw):
        try:
            return func(*args, **kw)
        except Exception as e:
            # print(e)
            pass
        return None

    def read(ole, fn):
        return ole.openstream(fn).read()

    ole, data, form = None, None, None
    try:
        ole = olefile.OleFileIO(oledata)
        try:
            _form = get_embedding_format(read(ole, ['\x01CompObj']))
            form = _form[2]
        except:
            pass
        # print(ole.listdir())
        for p in paths:
            data = try_default(read, ole, p)
            if data is not None: break
        return (form, data)
    except Exception as e:
        import sys
        print('failed to unpack ole:', e, file=sys.stderr)
    finally:
        if ole: ole.close()
    return (None, None)


def read_objs_from_doc(filename,*, paths=None, mapper=None, relpath=None, rIds=None):
    if mapper is None:
        mapper = lambda form, obj: (form, obj)
    rels = get_embedding_rels(filename, relpath)
    if rIds is None:
        ole_paths = [r.get('Target') for r in rels if r.get('Type', '').endswith('oleObject')]
    else:
        pathmap = dict(rels_to_paths(rels))
        ole_paths = [pathmap[r] for r in rIds]
    for ole in read_zip_files(filename, ole_paths, errors='ignore'):
        if ole is None: continue
        form, obj = read_ole_contents(ole, paths=paths)
        if obj is None:
            # print('empty', filename, form)
            continue
        # print(form, filename)
        try:
            _form, _obj = mapper(form, obj)
            if _obj is not None:
                yield (_form, _obj)
        except Exception as e:
            import sys
            print(filename, e, file=sys.stderr)
    return


def get_docx_table_embeddings(filename):
    from docx import Document
    from docx.oxml.shared import qn
    
    def get_table_row_rIds(row):
        ret = []
        for ps in [cell.paragraphs for cell in row.cells]:
            for rs in [p.runs for p in ps if p]:
                for x in [r.element.findall('{*}object/{*}OLEObject') for r in rs if r]:
                    if not x: continue
                    for o in x:
                        try:
                            ret.append(o.get(qn("r:id")))
                        except Exception as e: pass
        return ret
    
    doc = Document(filename)
    for tbl in doc.tables:
        ret = []
        row_rids = []
        header = None
        for n_row, row in enumerate(tbl.rows):
            vals = []
            for cell in row.cells:
                text = '\n'.join([p.text for p in cell.paragraphs if (p and p.text)])
                vals.append(text or None)
            ret.append(vals)
            row_rids.append(get_table_row_rIds(row))
        yield (ret, row_rids)
